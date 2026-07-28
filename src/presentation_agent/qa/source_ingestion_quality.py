"""Quality checks for local SourceDocument ingestion artifacts."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any


DEFAULT_SOURCE_DOCUMENT = Path("outputs/source_document.json")
DEFAULT_JSON_REPORT = Path("outputs/source_ingestion_quality_report.json")
DEFAULT_MD_REPORT = Path("outputs/source_ingestion_quality_report.md")

PAGE_LABEL_RE = re.compile(r"^Page\s+(\d+):", re.IGNORECASE)
PARAGRAPH_LABEL_RE = re.compile(r"^Paragraph\s+(\d+):", re.IGNORECASE)
FIGURE_RE = re.compile(r"\b(?:figure|fig\.)\s+[A-Za-z0-9.-]+", re.IGNORECASE)
TABLE_RE = re.compile(r"\btable\s+[A-Za-z0-9.-]+", re.IGNORECASE)


def build_source_ingestion_quality_report(
    *,
    source_document_path: str | Path,
    source_path: str | Path,
) -> dict[str, Any]:
    source_document_file = Path(source_document_path)
    original_source = Path(source_path)
    findings: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []

    source_document = _load_source_document(source_document_file, findings)
    source_type = str(source_document.get("source_type") or original_source.suffix.lower().lstrip(".") or "unknown")
    chunks = [chunk for chunk in source_document.get("chunks") or [] if isinstance(chunk, dict)]
    outline_items = [item for item in (source_document.get("outline") or {}).get("items") or [] if isinstance(item, dict)]
    figures = [item for item in source_document.get("figures") or [] if isinstance(item, dict)]
    tables = [item for item in source_document.get("tables") or [] if isinstance(item, dict)]
    citations = [item for item in source_document.get("citations") or [] if isinstance(item, dict)]
    evidence = [item for item in source_document.get("evidence") or [] if isinstance(item, dict)]
    text = "\n".join(str(chunk.get("text") or "") for chunk in chunks)

    source_file_metrics = _source_file_metrics(original_source, source_type, findings, warnings)
    extraction_metrics = _text_extraction_metrics(
        source_document=source_document,
        source_file_metrics=source_file_metrics,
        chunks=chunks,
        outline_items=outline_items,
        figures=figures,
        tables=tables,
        citations=citations,
        evidence=evidence,
        text=text,
    )
    structure = _structure_quality(
        source_type=source_type,
        chunks=chunks,
        outline_items=outline_items,
        tables=tables,
        figures=figures,
        citations=citations,
        evidence=evidence,
        text=text,
    )
    risk_flags = _risk_flags(
        source_type=source_type,
        extraction_metrics=extraction_metrics,
        structure=structure,
        figures=figures,
        tables=tables,
        citations=citations,
        source_warnings=source_document.get("warnings") or [],
    )
    for flag in risk_flags:
        severity = "severe" if flag["severity"] == "error" else "warning"
        target = findings if severity == "severe" else warnings
        target.append(_finding(flag["code"], severity, flag["message"], details=flag.get("details") or {}))

    severe_count = sum(1 for finding in findings if finding.get("severity") == "severe")
    warning_count = len(warnings)
    production_readiness = "failed" if severe_count else "warning" if warning_count else "passed"
    return {
        "schema_name": "source_ingestion_quality_report",
        "schema_version": "1.0",
        "status": production_readiness,
        "production_readiness": production_readiness,
        "source_path": _display_path(original_source),
        "source_document_path": _display_path(source_document_file),
        "source_type": source_type,
        "text_extraction": extraction_metrics,
        "structure_quality": structure,
        "risk_flags": risk_flags,
        "findings": findings + warnings,
        "findings_summary": {
            "total": len(findings) + warning_count,
            "severe": severe_count,
            "warning": warning_count,
        },
    }


def write_source_ingestion_quality_report(
    *,
    source_document_path: str | Path,
    source_path: str | Path,
    json_report_path: str | Path = DEFAULT_JSON_REPORT,
    md_report_path: str | Path = DEFAULT_MD_REPORT,
) -> dict[str, Any]:
    report = build_source_ingestion_quality_report(
        source_document_path=source_document_path,
        source_path=source_path,
    )
    json_path = Path(json_report_path)
    md_path = Path(md_report_path)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(report, indent=2, sort_keys=True, ensure_ascii=True) + "\n", encoding="utf-8")
    md_path.write_text(markdown_report(report), encoding="utf-8")
    return report


def markdown_report(report: dict[str, Any]) -> str:
    metrics = report.get("text_extraction") or {}
    structure = report.get("structure_quality") or {}
    lines = [
        "# Source Ingestion Quality Report",
        "",
        f"Status: `{report.get('status')}`",
        f"Source: `{report.get('source_path')}`",
        f"Source type: `{report.get('source_type')}`",
        f"Source artifact: `{report.get('source_document_path')}`",
        "",
        "## Text Extraction",
        "",
        f"- Character count: `{metrics.get('character_count')}`",
        f"- Page count: `{metrics.get('page_count')}`",
        f"- Empty page count: `{metrics.get('empty_page_count')}`",
        f"- Paragraph count: `{metrics.get('paragraph_count')}`",
        f"- Heading count: `{metrics.get('heading_count')}`",
        f"- Table count: `{metrics.get('table_count')}`",
        f"- Figure/caption candidate count: `{metrics.get('figure_caption_candidate_count')}`",
        f"- Citation/source anchor count: `{metrics.get('citation_source_anchor_count')}`",
        "",
        "## Structure Quality",
        "",
        f"- Heading hierarchy detected: `{structure.get('heading_hierarchy_detected')}`",
        f"- Section ordering: `{structure.get('section_ordering')}`",
        f"- Table extraction present: `{structure.get('table_extraction_present')}`",
        f"- Page anchors present: `{structure.get('page_anchors_present')}`",
        f"- Evidence anchors present: `{structure.get('evidence_anchors_present')}`",
        "",
        "## Risk Flags",
        "",
    ]
    flags = report.get("risk_flags") or []
    if flags:
        for flag in flags:
            lines.append(f"- `{flag.get('severity')}` `{flag.get('code')}`: {flag.get('message')}")
    else:
        lines.append("- None")
    lines.append("")
    return "\n".join(lines)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate source ingestion quality report for a SourceDocument artifact.")
    parser.add_argument("--source-document", type=Path, required=True)
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--json-report", type=Path, default=DEFAULT_JSON_REPORT)
    parser.add_argument("--md-report", type=Path, default=DEFAULT_MD_REPORT)
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    report = write_source_ingestion_quality_report(
        source_document_path=args.source_document,
        source_path=args.source,
        json_report_path=args.json_report,
        md_report_path=args.md_report,
    )
    print(f"WROTE {args.json_report}")
    print(f"SOURCE_INGESTION_QUALITY {report.get('status')}")
    return 1 if report.get("status") == "failed" else 0


def _load_source_document(path: Path, findings: list[dict[str, Any]]) -> dict[str, Any]:
    if not path.exists():
        findings.append(_finding("SOURCE_DOCUMENT_MISSING", "severe", f"SourceDocument artifact is missing: {_display_path(path)}"))
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001 - report malformed artifacts instead of crashing.
        findings.append(_finding("SOURCE_DOCUMENT_UNREADABLE", "severe", f"SourceDocument artifact could not be read: {type(exc).__name__}: {exc}"))
        return {}
    if payload.get("schema_name") != "source_document":
        findings.append(_finding("SOURCE_DOCUMENT_SCHEMA_UNEXPECTED", "severe", "Artifact does not declare schema_name=source_document."))
    return payload


def _source_file_metrics(source_path: Path, source_type: str, findings: list[dict[str, Any]], warnings: list[dict[str, Any]]) -> dict[str, Any]:
    metrics: dict[str, Any] = {"exists": source_path.exists(), "path": _display_path(source_path)}
    if not source_path.exists():
        findings.append(_finding("SOURCE_FILE_MISSING", "severe", f"Original source file is missing: {_display_path(source_path)}"))
        return metrics
    metrics["size_bytes"] = source_path.stat().st_size
    if source_type == "pdf":
        metrics.update(_pdf_file_metrics(source_path, warnings))
    elif source_type == "docx":
        metrics.update(_docx_file_metrics(source_path, warnings))
    return metrics


def _pdf_file_metrics(source_path: Path, warnings: list[dict[str, Any]]) -> dict[str, Any]:
    try:
        import fitz  # type: ignore
    except ImportError:
        warnings.append(_finding("PDF_PAGE_COUNT_UNAVAILABLE", "warning", "PyMuPDF is unavailable; page count was inferred from extracted page anchors."))
        return {}
    try:
        document = fitz.open(source_path)
        try:
            page_count = int(document.page_count)
            empty_pages = 0
            for page in document:
                if not str(page.get_text("text") or "").strip():
                    empty_pages += 1
            return {"original_page_count": page_count, "original_empty_page_count": empty_pages}
        finally:
            document.close()
    except Exception as exc:  # noqa: BLE001 - optional source metrics should be warnings.
        warnings.append(_finding("PDF_PAGE_COUNT_UNAVAILABLE", "warning", f"PDF page metrics could not be computed: {type(exc).__name__}: {exc}"))
        return {}


def _docx_file_metrics(source_path: Path, warnings: list[dict[str, Any]]) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        warnings.append(_finding("DOCX_STYLE_METRICS_UNAVAILABLE", "warning", "python-docx is unavailable; DOCX style metrics were inferred from SourceDocument only."))
        return {}
    try:
        document = Document(source_path)
    except Exception as exc:  # noqa: BLE001
        warnings.append(_finding("DOCX_STYLE_METRICS_UNAVAILABLE", "warning", f"DOCX style metrics could not be computed: {type(exc).__name__}: {exc}"))
        return {}
    heading_style_count = 0
    paragraph_count = 0
    inline_shape_count = len(getattr(document, "inline_shapes", []) or [])
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            paragraph_count += 1
            if str(paragraph.style.name or "").lower().startswith("heading"):
                heading_style_count += 1
    return {
        "original_paragraph_count": paragraph_count,
        "original_heading_style_count": heading_style_count,
        "original_table_count": len(document.tables),
        "original_inline_shape_count": inline_shape_count,
    }


def _text_extraction_metrics(
    *,
    source_document: dict[str, Any],
    source_file_metrics: dict[str, Any],
    chunks: list[dict[str, Any]],
    outline_items: list[dict[str, Any]],
    figures: list[dict[str, Any]],
    tables: list[dict[str, Any]],
    citations: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
    text: str,
) -> dict[str, Any]:
    page_numbers = sorted(
        {
            int(match.group(1))
            for chunk in chunks
            if (match := PAGE_LABEL_RE.match(str(chunk.get("text") or "")))
        }
    )
    paragraph_numbers = sorted(
        {
            int(match.group(1))
            for chunk in chunks
            if (match := PARAGRAPH_LABEL_RE.match(str(chunk.get("text") or "")))
        }
    )
    page_count = source_file_metrics.get("original_page_count") or (max(page_numbers) if page_numbers else None)
    original_empty_pages = source_file_metrics.get("original_empty_page_count")
    extracted_pages = len(page_numbers)
    empty_page_count = original_empty_pages if original_empty_pages is not None else max(0, int(page_count or 0) - extracted_pages)
    figure_caption_candidates = len(figures) + len(FIGURE_RE.findall(text))
    table_caption_candidates = len(tables) + len(TABLE_RE.findall(text))
    return {
        "character_count": len(text),
        "chunk_count": len(chunks),
        "page_count": page_count,
        "extracted_page_anchor_count": extracted_pages,
        "empty_page_count": empty_page_count,
        "paragraph_count": len(paragraph_numbers) if paragraph_numbers else len(chunks),
        "heading_count": len(outline_items),
        "heading_level_count": len({item.get("level") for item in outline_items if item.get("level") is not None}),
        "table_count": len(tables),
        "figure_count": len(figures),
        "figure_caption_candidate_count": figure_caption_candidates,
        "table_caption_candidate_count": table_caption_candidates,
        "citation_source_anchor_count": len(citations),
        "evidence_anchor_count": len(evidence),
        "source_file_metrics": source_file_metrics,
    }


def _structure_quality(
    *,
    source_type: str,
    chunks: list[dict[str, Any]],
    outline_items: list[dict[str, Any]],
    tables: list[dict[str, Any]],
    figures: list[dict[str, Any]],
    citations: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
    text: str,
) -> dict[str, Any]:
    levels = [int(item.get("level")) for item in outline_items if isinstance(item.get("level"), int)]
    source_chunk_ids = [str(chunk.get("chunk_id") or "") for chunk in chunks]
    page_anchor_labels = [citation for citation in citations if str(citation.get("label") or "").lower().startswith("page ")]
    citation_chunk_ids = {chunk_id for citation in citations for chunk_id in citation.get("source_chunk_ids") or []}
    evidence_chunk_ids = {str(item.get("source_chunk_id") or "") for item in evidence}
    return {
        "source_outline_quality": "usable" if outline_items else "weak",
        "heading_hierarchy_detected": bool(outline_items),
        "heading_levels": sorted(set(levels)),
        "heading_hierarchy_depth": len(set(levels)),
        "section_ordering": "ordered" if _outline_ordered(outline_items, source_chunk_ids) else "not_detected",
        "table_extraction_present": bool(tables),
        "figure_caption_extraction_present": bool(figures),
        "page_anchors_present": bool(page_anchor_labels) if source_type == "pdf" else None,
        "source_anchors_present": bool(citation_chunk_ids),
        "evidence_anchors_present": bool(evidence_chunk_ids),
        "all_evidence_anchors_resolve": evidence_chunk_ids.issubset(set(source_chunk_ids)),
        "captions_anchored": all(item.get("source_chunk_id") for item in figures + tables),
        "detected_table_caption_text": bool(TABLE_RE.search(text)),
        "detected_figure_caption_text": bool(FIGURE_RE.search(text)),
    }


def _risk_flags(
    *,
    source_type: str,
    extraction_metrics: dict[str, Any],
    structure: dict[str, Any],
    figures: list[dict[str, Any]],
    tables: list[dict[str, Any]],
    citations: list[dict[str, Any]],
    source_warnings: list[str],
) -> list[dict[str, Any]]:
    flags: list[dict[str, Any]] = []
    char_count = int(extraction_metrics.get("character_count") or 0)
    page_count = extraction_metrics.get("page_count")
    empty_pages = int(extraction_metrics.get("empty_page_count") or 0)
    if source_type == "pdf" and page_count and char_count < max(120, int(page_count) * 80):
        flags.append(_risk("SCANNED_PDF_LIKELY", "warning", "PDF text extraction is low relative to page count; scanned or image-only content is possible."))
    if source_type == "pdf" and page_count and empty_pages >= int(page_count):
        flags.append(_risk("SCANNED_PDF_LIKELY", "error", "No extractable text was detected on PDF pages."))
    if char_count < 80:
        flags.append(_risk("TOO_LITTLE_TEXT_EXTRACTED", "error", "Too little text was extracted for reliable presentation planning.", {"character_count": char_count}))
    elif char_count < 240:
        flags.append(_risk("LOW_TEXT_VOLUME", "warning", "Extracted text is sparse; resulting plans may be shallow.", {"character_count": char_count}))
    if not structure.get("heading_hierarchy_detected"):
        flags.append(_risk("NO_HEADINGS_DETECTED", "warning", "No heading hierarchy was detected."))
    if source_type == "pdf" and not structure.get("page_anchors_present"):
        flags.append(_risk("NO_PAGE_ANCHORS", "warning", "PDF ingestion did not preserve page-level source anchors."))
    if structure.get("detected_table_caption_text") and not tables:
        flags.append(_risk("TABLES_DETECTED_NOT_REPRESENTED", "warning", "Table caption text was detected but no table references were represented."))
    if (figures or tables) and not structure.get("captions_anchored"):
        flags.append(_risk("CAPTIONS_NOT_ANCHORED", "warning", "One or more figure/table captions lack source chunk anchors."))
    if source_type == "docx":
        level_count = int(extraction_metrics.get("heading_level_count") or 0)
        if not structure.get("heading_hierarchy_detected"):
            flags.append(_risk("DOCX_STYLES_MISSING", "warning", "DOCX headings were not detected from paragraph styles."))
        elif len(citations) > 0 and level_count <= 1 and extraction_metrics.get("heading_count", 0) > 3:
            flags.append(_risk("DOCX_STYLES_FLAT", "warning", "DOCX heading styles appear flat; section hierarchy may need review."))
    for warning in source_warnings:
        if warning == "source_document_empty":
            flags.append(_risk("SOURCE_DOCUMENT_EMPTY", "error", "Ingestion reported an empty source document."))
        elif warning == "pdf_page_without_extractable_text":
            flags.append(_risk("PDF_PAGE_WITHOUT_EXTRACTABLE_TEXT", "warning", "At least one PDF page had no extractable text."))
        elif warning == "insufficient_source_structure":
            flags.append(_risk("INSUFFICIENT_SOURCE_STRUCTURE", "warning", "Ingestion reported weak or missing source structure."))
    return _dedupe_risks(flags)


def _outline_ordered(outline_items: list[dict[str, Any]], source_chunk_ids: list[str]) -> bool:
    if not outline_items:
        return False
    order = {chunk_id: index for index, chunk_id in enumerate(source_chunk_ids)}
    positions: list[int] = []
    for item in outline_items:
        ids = [chunk_id for chunk_id in item.get("source_chunk_ids") or [] if chunk_id in order]
        if ids:
            positions.append(min(order[chunk_id] for chunk_id in ids))
    return bool(positions) and positions == sorted(positions)


def _risk(code: str, severity: str, message: str, details: dict[str, Any] | None = None) -> dict[str, Any]:
    payload = {"code": code, "severity": severity, "message": message}
    if details:
        payload["details"] = details
    return payload


def _finding(code: str, severity: str, message: str, *, details: dict[str, Any] | None = None) -> dict[str, Any]:
    payload = {"code": code, "severity": severity, "message": message}
    if details:
        payload["details"] = details
    return payload


def _dedupe_risks(flags: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[str] = set()
    result: list[dict[str, Any]] = []
    for flag in flags:
        code = str(flag.get("code") or "")
        if code in seen:
            continue
        seen.add(code)
        result.append(flag)
    return result


def _display_path(path: str | Path) -> str:
    return Path(path).as_posix()


if __name__ == "__main__":
    raise SystemExit(main())
