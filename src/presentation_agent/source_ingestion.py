"""Deterministic local source ingestion for source-to-deck planning."""

from __future__ import annotations

import re
from pathlib import Path
from statistics import median
from typing import Any

from .source_planning import (
    SourceCitation,
    SourceChunk,
    SourceDocument,
    SourceEvidence,
    SourceFigureRef,
    SourceOutline,
    SourceOutlineItem,
    SourceTableRef,
    with_structural_hash,
    write_source_planning_json,
)


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
FIGURE_RE = re.compile(r"\b(?:figure|fig\.)\s+([A-Za-z0-9.-]+)", re.IGNORECASE)
TABLE_RE = re.compile(r"\btable\s+([A-Za-z0-9.-]+)", re.IGNORECASE)
FIGURE_CAPTION_RE = re.compile(r"^\s*(Figure|Fig\.)\s+([A-Za-z0-9.-]+)\s*[:.\-]?\s*(.+)$", re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(r"^\s*Table\s+([A-Za-z0-9.-]+)\s*[:.\-]?\s*(.+)$", re.IGNORECASE)
CITATION_RE = re.compile(r"(\[[0-9,\s-]{1,16}\]|\([A-Z][A-Za-z-]+(?:\s+et\s+al\.)?,\s*\d{4}[a-z]?\)|doi:\s*\S+)", re.IGNORECASE)
MAX_CHUNK_CHARS = 1400


def ingest_source_file(path: str | Path) -> SourceDocument:
    source_path = Path(path).resolve()
    if not source_path.is_file():
        raise FileNotFoundError(f"source file not found: {source_path}")
    suffix = source_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = source_path.read_text(encoding="utf-8")
        source_type = "md" if suffix == ".md" else "txt"
        document = _ingest_text(
            text=text,
            source_path=source_path,
            source_type=source_type,
        )
    elif suffix == ".pdf":
        document = _ingest_pdf(source_path)
    elif suffix == ".docx":
        document = _ingest_docx(source_path)
    else:
        raise ValueError("source ingestion supports .txt, .md, .pdf, and .docx")
    return with_structural_hash(document)  # type: ignore[return-value]


def write_source_document(document: SourceDocument, output_path: str | Path) -> Path:
    return write_source_planning_json(document, output_path)


def _ingest_pdf(source_path: Path) -> SourceDocument:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PDF ingestion requires PyMuPDF (`fitz`). Install pymupdf to ingest .pdf sources.") from exc

    pages: list[dict[str, Any]] = []
    try:
        document = fitz.open(source_path)
    except Exception as exc:
        raise ValueError(f"could not open PDF source {source_path}: {exc}") from exc

    try:
        for page_index, page in enumerate(document, start=1):
            pages.append(_extract_pdf_page(page, page_index))
    finally:
        document.close()

    chunks: list[SourceChunk] = []
    outline_entries: list[SourceOutlineItem] = []
    heading_stack: list[tuple[int, str]] = []
    current_heading_path: list[str] = []
    current_heading_level: int | None = None
    char_offset = 0
    title: str | None = None

    for page in pages:
        for block in page["blocks"]:
            text = _clean_text(block["text"])
            if not text:
                continue
            if block["is_heading"]:
                level = int(block["level"])
                title = title or text
                heading_stack = [(lvl, value) for lvl, value in heading_stack if lvl < level]
                heading_stack.append((level, text))
                current_heading_path = [value for _, value in heading_stack]
                current_heading_level = level
                outline_entries.append(
                    SourceOutlineItem(
                        outline_id=f"outline-{len(outline_entries) + 1:03d}",
                        title=text,
                        level=level,
                        source_chunk_ids=[],
                    )
                )
                continue
            chunk_text = f"Page {page['page_number']}: {text}"
            start_char = char_offset
            end_char = start_char + len(chunk_text)
            chunk_id = f"chunk-{len(chunks) + 1:03d}"
            chunks.append(
                SourceChunk(
                    chunk_id=chunk_id,
                    text=chunk_text,
                    heading_path=list(current_heading_path),
                    heading_level=current_heading_level,
                    start_line=int(page["page_number"]),
                    end_line=int(page["page_number"]),
                    start_char=start_char,
                    end_char=end_char,
                    token_estimate=_token_estimate(chunk_text),
                )
            )
            char_offset = end_char + 1

    _attach_outline_chunks(outline_entries, chunks)
    figures, tables = _extract_refs(chunks)
    _merge_caption_refs(figures, tables, chunks)
    evidence = _extract_evidence(chunks)
    citations = _extract_citations(chunks)
    citations.extend(_page_citations(chunks, "Page"))
    warnings: list[str] = []
    if not outline_entries:
        warnings.append("insufficient_source_structure")
    if not chunks:
        warnings.append("source_document_empty")
    if any(page["text_empty"] for page in pages):
        warnings.append("pdf_page_without_extractable_text")
    return SourceDocument(
        document_id=_slug(source_path.stem) or "source-document",
        source_path=str(source_path),
        source_type="pdf",
        title=title or _title_from_text([block["text"] for page in pages for block in page["blocks"]]) or source_path.stem,
        chunks=chunks,
        outline=SourceOutline(
            items=outline_entries,
            structure_quality=_structure_quality(outline_entries, chunks),
        ),
        evidence=evidence,
        figures=figures,
        tables=tables,
        citations=_dedupe_citations(citations),
        warnings=warnings,
    )


def _extract_pdf_page(page: Any, page_number: int) -> dict[str, Any]:
    payload = page.get_text("dict")
    raw_blocks = []
    sizes: list[float] = []
    for block in payload.get("blocks", []):
        if block.get("type") != 0:
            continue
        lines = []
        max_size = 0.0
        bold = False
        for line in block.get("lines", []):
            line_text_parts = []
            for span in line.get("spans", []):
                text = str(span.get("text") or "")
                if text.strip():
                    line_text_parts.append(text)
                size = float(span.get("size") or 0)
                if size:
                    sizes.append(size)
                    max_size = max(max_size, size)
                font = str(span.get("font") or "")
                flags = int(span.get("flags") or 0)
                bold = bold or "bold" in font.lower() or bool(flags & 16)
            line_text = _clean_text(" ".join(line_text_parts))
            if line_text:
                lines.append(line_text)
        text = _clean_text("\n".join(lines))
        if text:
            raw_blocks.append({"text": text, "max_size": max_size, "bold": bold})
    median_size = median(sizes) if sizes else 0.0
    blocks = []
    for block in raw_blocks:
        first_line = block["text"].splitlines()[0].strip()
        is_heading = _pdf_heading_candidate(first_line, block["max_size"], block["bold"], median_size)
        blocks.append(
            {
                "text": block["text"],
                "is_heading": is_heading,
                "level": 1 if block["max_size"] >= median_size + 5 else 2,
            }
        )
    return {"page_number": page_number, "blocks": blocks, "text_empty": not bool(raw_blocks)}


def _pdf_heading_candidate(text: str, font_size: float, bold: bool, median_size: float) -> bool:
    if not text or len(text) > 120 or text.endswith("."):
        return False
    if FIGURE_CAPTION_RE.match(text) or TABLE_CAPTION_RE.match(text):
        return False
    if font_size >= median_size + 2.0:
        return True
    return bold and len(text.split()) <= 12 and font_size >= median_size + 0.5


def _ingest_docx(source_path: Path) -> SourceDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX ingestion requires python-docx. Install python-docx to ingest .docx sources.") from exc

    try:
        document = Document(source_path)
    except Exception as exc:
        raise ValueError(f"could not open DOCX source {source_path}: {exc}") from exc

    chunks: list[SourceChunk] = []
    outline_entries: list[SourceOutlineItem] = []
    figures: list[SourceFigureRef] = []
    tables: list[SourceTableRef] = []
    heading_stack: list[tuple[int, str]] = []
    current_heading_path: list[str] = []
    current_heading_level: int | None = None
    char_offset = 0
    title: str | None = None

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = _clean_text(paragraph.text)
        if not text:
            continue
        heading_level = _docx_heading_level(paragraph)
        if heading_level is not None:
            title = title or text
            heading_stack = [(lvl, value) for lvl, value in heading_stack if lvl < heading_level]
            heading_stack.append((heading_level, text))
            current_heading_path = [value for _, value in heading_stack]
            current_heading_level = heading_level
            outline_entries.append(
                SourceOutlineItem(
                    outline_id=f"outline-{len(outline_entries) + 1:03d}",
                    title=text,
                    level=heading_level,
                    source_chunk_ids=[],
                )
            )
            continue
        chunk_text = f"Paragraph {paragraph_index}: {text}"
        chunk, char_offset = _make_chunk(
            chunk_id=f"chunk-{len(chunks) + 1:03d}",
            text=chunk_text,
            heading_path=current_heading_path,
            heading_level=current_heading_level,
            start_line=paragraph_index,
            end_line=paragraph_index,
            start_char=char_offset,
        )
        chunks.append(chunk)
        if _paragraph_has_inline_shape(paragraph):
            figures.append(
                SourceFigureRef(
                    figure_id=f"inline-shape-{len(figures) + 1:03d}",
                    label=f"Inline shape near paragraph {paragraph_index}",
                    caption=None,
                    source_chunk_id=chunk.chunk_id,
                )
            )

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if not rows:
            continue
        table_text = f"Table {table_index}: " + "\n".join(rows)
        chunk, char_offset = _make_chunk(
            chunk_id=f"chunk-{len(chunks) + 1:03d}",
            text=table_text,
            heading_path=current_heading_path,
            heading_level=current_heading_level,
            start_line=len(document.paragraphs) + table_index,
            end_line=len(document.paragraphs) + table_index,
            start_char=char_offset,
        )
        chunks.append(chunk)
        tables.append(
            SourceTableRef(
                table_id=f"table-{table_index}",
                label=f"Table {table_index}",
                caption=rows[0][:220] if rows else None,
                source_chunk_id=chunk.chunk_id,
            )
        )

    _attach_outline_chunks(outline_entries, chunks)
    extracted_figures, extracted_tables = _extract_refs(chunks)
    figures = _dedupe_figures([*figures, *extracted_figures])
    tables = _dedupe_tables([*tables, *extracted_tables])
    _merge_caption_refs(figures, tables, chunks)
    citations = _dedupe_citations([*_extract_citations(chunks), *_paragraph_citations(chunks)])
    warnings: list[str] = []
    if not outline_entries:
        warnings.append("insufficient_source_structure")
    if not chunks:
        warnings.append("source_document_empty")
    return SourceDocument(
        document_id=_slug(source_path.stem) or "source-document",
        source_path=str(source_path),
        source_type="docx",
        title=title or _title_from_text([paragraph.text for paragraph in document.paragraphs]) or source_path.stem,
        chunks=chunks,
        outline=SourceOutline(
            items=outline_entries,
            structure_quality=_structure_quality(outline_entries, chunks),
        ),
        evidence=_extract_evidence(chunks),
        figures=figures,
        tables=tables,
        citations=citations,
        warnings=warnings,
    )


def _ingest_text(*, text: str, source_path: Path, source_type: str) -> SourceDocument:
    lines = text.splitlines()
    line_start_offsets = _line_start_offsets(text)
    chunks: list[SourceChunk] = []
    outline_entries: list[SourceOutlineItem] = []
    heading_stack: list[tuple[int, str]] = []
    current_heading_path: list[str] = []
    current_heading_level: int | None = None
    current_start_line = 1
    current_lines: list[tuple[int, str]] = []
    title: str | None = None

    def flush() -> None:
        nonlocal current_lines, current_start_line
        if not current_lines:
            return
        for piece in _split_chunk_lines(current_lines):
            start_line = piece[0][0]
            end_line = piece[-1][0]
            chunk_text = "\n".join(line for _, line in piece).strip()
            if not chunk_text:
                continue
            start_char = line_start_offsets[start_line - 1]
            end_char = line_start_offsets[end_line - 1] + len(lines[end_line - 1])
            chunk_id = f"chunk-{len(chunks) + 1:03d}"
            chunks.append(
                SourceChunk(
                    chunk_id=chunk_id,
                    text=chunk_text,
                    heading_path=list(current_heading_path),
                    heading_level=current_heading_level,
                    start_line=start_line,
                    end_line=end_line,
                    start_char=start_char,
                    end_char=end_char,
                    token_estimate=_token_estimate(chunk_text),
                )
            )
        current_lines = []
        current_start_line = end_line + 1 if "end_line" in locals() else current_start_line

    for line_no, line in enumerate(lines, start=1):
        match = HEADING_RE.match(line) if source_type == "md" else None
        if match:
            flush()
            level = len(match.group(1))
            heading = match.group(2).strip()
            title = title or heading
            heading_stack = [(lvl, text) for lvl, text in heading_stack if lvl < level]
            heading_stack.append((level, heading))
            current_heading_path = [text for _, text in heading_stack]
            current_heading_level = level
            outline_entries.append(
                SourceOutlineItem(
                    outline_id=f"outline-{len(outline_entries) + 1:03d}",
                    title=heading,
                    level=level,
                    source_chunk_ids=[],
                )
            )
            current_start_line = line_no + 1
            continue
        current_lines.append((line_no, line))
    flush()

    _attach_outline_chunks(outline_entries, chunks)
    figures, tables = _extract_refs(chunks)
    evidence = _extract_evidence(chunks)
    warnings: list[str] = []
    if not outline_entries:
        warnings.append("insufficient_source_structure")
    if not chunks:
        warnings.append("source_document_empty")
    document_id = _slug(source_path.stem) or "source-document"
    return SourceDocument(
        document_id=document_id,
        source_path=str(source_path),
        source_type=source_type,  # type: ignore[arg-type]
        title=title or _title_from_text(lines) or source_path.stem,
        chunks=chunks,
        outline=SourceOutline(
            items=outline_entries,
            structure_quality=_structure_quality(outline_entries, chunks),
        ),
        evidence=evidence,
        figures=figures,
        tables=tables,
        warnings=warnings,
    )


def _line_start_offsets(text: str) -> list[int]:
    offsets = [0]
    for index, char in enumerate(text):
        if char == "\n":
            offsets.append(index + 1)
    return offsets


def _split_chunk_lines(lines: list[tuple[int, str]]) -> list[list[tuple[int, str]]]:
    pieces: list[list[tuple[int, str]]] = []
    current: list[tuple[int, str]] = []
    current_chars = 0
    for item in lines:
        line_len = len(item[1]) + 1
        if current and current_chars + line_len > MAX_CHUNK_CHARS:
            pieces.append(current)
            current = []
            current_chars = 0
        current.append(item)
        current_chars += line_len
    if current:
        pieces.append(current)
    return pieces


def _attach_outline_chunks(outline_entries: list[SourceOutlineItem], chunks: list[SourceChunk]) -> None:
    for entry in outline_entries:
        entry.source_chunk_ids.clear()
        for chunk in chunks:
            if chunk.heading_path and chunk.heading_path[-1] == entry.title:
                entry.source_chunk_ids.append(chunk.chunk_id)


def _extract_refs(chunks: list[SourceChunk]) -> tuple[list[SourceFigureRef], list[SourceTableRef]]:
    figures: list[SourceFigureRef] = []
    tables: list[SourceTableRef] = []
    seen_figures: set[str] = set()
    seen_tables: set[str] = set()
    for chunk in chunks:
        for match in FIGURE_RE.finditer(chunk.text):
            label = f"Figure {match.group(1)}"
            figure_id = _slug(label) or f"figure-{len(figures) + 1:03d}"
            if figure_id not in seen_figures:
                seen_figures.add(figure_id)
                figures.append(SourceFigureRef(figure_id=figure_id, label=label, source_chunk_id=chunk.chunk_id))
        for match in TABLE_RE.finditer(chunk.text):
            label = f"Table {match.group(1)}"
            table_id = _slug(label) or f"table-{len(tables) + 1:03d}"
            if table_id not in seen_tables:
                seen_tables.add(table_id)
                tables.append(SourceTableRef(table_id=table_id, label=label, source_chunk_id=chunk.chunk_id))
    return figures, tables


def _extract_evidence(chunks: list[SourceChunk]) -> list[SourceEvidence]:
    evidence: list[SourceEvidence] = []
    for chunk in chunks:
        first_sentence = _first_sentence(chunk.text)
        if first_sentence:
            evidence.append(
                SourceEvidence(
                    evidence_id=f"evidence-{len(evidence) + 1:03d}",
                    source_chunk_id=chunk.chunk_id,
                    quote=first_sentence[:280],
                    relevance="first salient source sentence in chunk",
                    confidence=0.65,
                )
            )
    return evidence


def _clean_text(value: str) -> str:
    return " ".join(str(value).replace("\x00", " ").split())


def _make_chunk(
    *,
    chunk_id: str,
    text: str,
    heading_path: list[str],
    heading_level: int | None,
    start_line: int,
    end_line: int,
    start_char: int,
) -> tuple[SourceChunk, int]:
    end_char = start_char + len(text)
    return (
        SourceChunk(
            chunk_id=chunk_id,
            text=text,
            heading_path=list(heading_path),
            heading_level=heading_level,
            start_line=start_line,
            end_line=end_line,
            start_char=start_char,
            end_char=end_char,
            token_estimate=_token_estimate(text),
        ),
        end_char + 1,
    )


def _docx_heading_level(paragraph: Any) -> int | None:
    style_name = ""
    if getattr(paragraph, "style", None) is not None:
        style_name = str(getattr(paragraph.style, "name", "") or "")
    normalized = style_name.strip().lower()
    if normalized == "title":
        return 1
    match = re.match(r"heading\s+([1-6])\b", normalized)
    if match:
        return int(match.group(1))
    return None


def _paragraph_has_inline_shape(paragraph: Any) -> bool:
    xml = getattr(getattr(paragraph, "_p", None), "xml", "")
    return "<w:drawing" in xml or "<w:pict" in xml


def _merge_caption_refs(
    figures: list[SourceFigureRef],
    tables: list[SourceTableRef],
    chunks: list[SourceChunk],
) -> None:
    figure_by_label = {figure.label.lower(): figure for figure in figures}
    table_by_label = {table.label.lower(): table for table in tables}
    for chunk in chunks:
        text = _strip_source_anchor_prefix(chunk.text)
        figure_match = FIGURE_CAPTION_RE.match(text)
        if figure_match:
            label = f"Figure {figure_match.group(2)}"
            caption = figure_match.group(3).strip() or None
            existing = figure_by_label.get(label.lower())
            if existing:
                existing.caption = existing.caption or caption
                existing.source_chunk_id = existing.source_chunk_id or chunk.chunk_id
            else:
                figure = SourceFigureRef(
                    figure_id=_slug(label) or f"figure-{len(figures) + 1:03d}",
                    label=label,
                    caption=caption,
                    source_chunk_id=chunk.chunk_id,
                )
                figures.append(figure)
                figure_by_label[label.lower()] = figure
        table_match = TABLE_CAPTION_RE.match(text)
        if table_match:
            label = f"Table {table_match.group(1)}"
            caption = table_match.group(2).strip() or None
            existing = table_by_label.get(label.lower())
            if existing:
                existing.caption = existing.caption or caption
                existing.source_chunk_id = existing.source_chunk_id or chunk.chunk_id
            else:
                table = SourceTableRef(
                    table_id=_slug(label) or f"table-{len(tables) + 1:03d}",
                    label=label,
                    caption=caption,
                    source_chunk_id=chunk.chunk_id,
                )
                tables.append(table)
                table_by_label[label.lower()] = table


def _strip_source_anchor_prefix(text: str) -> str:
    return re.sub(r"^(?:Page|Paragraph)\s+\d+\s*:\s*", "", text.strip())


def _extract_citations(chunks: list[SourceChunk]) -> list[SourceCitation]:
    citations: list[SourceCitation] = []
    seen: dict[str, SourceCitation] = {}
    for chunk in chunks:
        for match in CITATION_RE.finditer(chunk.text):
            label = match.group(1).strip()
            citation_id = _slug(label) or f"citation-{len(citations) + 1:03d}"
            existing = seen.get(citation_id)
            if existing:
                if chunk.chunk_id not in existing.source_chunk_ids:
                    existing.source_chunk_ids.append(chunk.chunk_id)
                continue
            citation = SourceCitation(citation_id=citation_id, label=label, source_chunk_ids=[chunk.chunk_id])
            seen[citation_id] = citation
            citations.append(citation)
    return citations


def _page_citations(chunks: list[SourceChunk], label_prefix: str) -> list[SourceCitation]:
    grouped: dict[int, list[str]] = {}
    for chunk in chunks:
        grouped.setdefault(chunk.start_line, []).append(chunk.chunk_id)
    return [
        SourceCitation(
            citation_id=f"{_slug(label_prefix)}-{page_number:03d}",
            label=f"{label_prefix} {page_number}",
            source_chunk_ids=chunk_ids,
        )
        for page_number, chunk_ids in sorted(grouped.items())
    ]


def _paragraph_citations(chunks: list[SourceChunk]) -> list[SourceCitation]:
    return [
        SourceCitation(
            citation_id=f"source-anchor-{chunk.chunk_id}",
            label=f"Source anchor {chunk.chunk_id}",
            source_chunk_ids=[chunk.chunk_id],
        )
        for chunk in chunks
    ]


def _dedupe_citations(citations: list[SourceCitation]) -> list[SourceCitation]:
    merged: dict[str, SourceCitation] = {}
    for citation in citations:
        key = citation.citation_id
        existing = merged.get(key)
        if existing is None:
            merged[key] = citation
            continue
        for chunk_id in citation.source_chunk_ids:
            if chunk_id not in existing.source_chunk_ids:
                existing.source_chunk_ids.append(chunk_id)
    return list(merged.values())


def _dedupe_figures(figures: list[SourceFigureRef]) -> list[SourceFigureRef]:
    merged: dict[str, SourceFigureRef] = {}
    for figure in figures:
        key = figure.figure_id
        existing = merged.get(key)
        if existing is None:
            merged[key] = figure
            continue
        existing.caption = existing.caption or figure.caption
        existing.source_chunk_id = existing.source_chunk_id or figure.source_chunk_id
    return list(merged.values())


def _dedupe_tables(tables: list[SourceTableRef]) -> list[SourceTableRef]:
    merged: dict[str, SourceTableRef] = {}
    for table in tables:
        key = table.table_id
        existing = merged.get(key)
        if existing is None:
            merged[key] = table
            continue
        existing.caption = existing.caption or table.caption
        existing.source_chunk_id = existing.source_chunk_id or table.source_chunk_id
    return list(merged.values())


def _first_sentence(text: str) -> str:
    compact = " ".join(text.split())
    for delimiter in (". ", "? ", "! "):
        if delimiter in compact:
            return compact.split(delimiter, 1)[0].strip() + delimiter.strip()
    return compact[:240].strip()


def _title_from_text(lines: list[str]) -> str | None:
    for line in lines:
        text = line.strip()
        if text:
            return text[:80]
    return None


def _structure_quality(outline_entries: list[SourceOutlineItem], chunks: list[SourceChunk]) -> str:
    if not outline_entries:
        return "weak" if chunks else "none"
    if len(outline_entries) >= 4:
        return "strong"
    return "usable"


def _token_estimate(text: str) -> int:
    return max(1, len(text.split()))


def _slug(value: str) -> str:
    words = "".join(char.lower() if char.isalnum() else " " for char in value).split()
    return "-".join(words)
